"""
Document processing module for handling various file types
"""
import os
import PyPDF2
import docx
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path
import hashlib

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles processing of various document types (PDF, TXT, DOCX)
    """
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.txt', '.docx', '.doc', '.jpg', '.jpeg', '.png'}
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract text content
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Dict[str, Any]: Document metadata and content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        # Generate file hash for unique identification
        file_hash = self._generate_file_hash(file_path)
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            text_content = self._extract_pdf_text(file_path)
        elif file_path.suffix.lower() == '.txt':
            text_content = self._extract_txt_text(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            text_content = self._extract_docx_text(file_path)
        elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
            text_content = self._extract_image_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        # Clean and preprocess text
        cleaned_text = self._clean_text(text_content)
        
        return {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_hash': file_hash,
            'file_type': file_path.suffix.lower(),
            'content': cleaned_text,
            'word_count': len(cleaned_text.split()),
            'char_count': len(cleaned_text)
        }
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text() or ""
                        if not page_text.strip():
                            # Fallback to OCR for scanned pages if available
                            try:
                                from pdf2image import convert_from_path
                                from PIL import Image
                                import pytesseract
                                images = convert_from_path(str(file_path), first_page=page_num + 1, last_page=page_num + 1)
                                if images:
                                    page_text = pytesseract.image_to_string(images[0]) or ""
                            except Exception as ocr_err:
                                logger.warning(f"OCR fallback failed on PDF page {page_num + 1}: {ocr_err}")
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
                
            return text_content
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise

    def _extract_image_text(self, file_path: Path) -> str:
        """Extract text from an image using OCR if available"""
        try:
            try:
                from PIL import Image
                import pytesseract
            except Exception as import_err:
                logger.warning(f"OCR dependencies not available for image extraction: {import_err}")
                return ""
            image = Image.open(file_path)
            return pytesseract.image_to_string(image) or ""
        except Exception as e:
            logger.error(f"Error processing image file {file_path}: {str(e)}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file {file_path}: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error processing TXT file {file_path}: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        if not text:
            return ""
        
        # Normalize whitespace but keep page separators for page mapping
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            cleaned_lines.append(line)
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def _generate_file_hash(self, file_path: Path) -> str:
        """Generate SHA-256 hash of file for unique identification"""
        hash_sha256 = hashlib.sha256()
        
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        
        return hash_sha256.hexdigest()
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks for better processing
        
        Args:
            text (str): Text to chunk
            chunk_size (int): Size of each chunk
            overlap (int): Overlap between chunks
            
        Returns:
            List[Dict[str, Any]]: List of text chunks with metadata
        """
        if not text:
            return []
        
        chunks = []

        # Build page marker index to map positions to page numbers
        page_markers: List[Tuple[int, int]] = []  # (position_index, page_number)
        import re as _re
        for match in _re.finditer(r"---\s*Page\s*(\d+)\s*---", text):
            try:
                page_num = int(match.group(1))
                page_markers.append((match.start(), page_num))
            except Exception:
                continue

        def find_page_for_pos(pos: int) -> Optional[int]:
            if not page_markers:
                return None
            current = None
            for marker_pos, pnum in page_markers:
                if marker_pos <= pos:
                    current = pnum
                else:
                    break
            return current
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:  # Only add non-empty chunks
                chunks.append({
                    'chunk_id': chunk_id,
                    'text': chunk_text,
                    'start_pos': start,
                    'end_pos': end,
                    'length': len(chunk_text),
                    'page_number': find_page_for_pos(start)
                })
                chunk_id += 1
            
            # Move start position with overlap
            start = end - overlap
            if start <= 0:
                start = end
        
        return chunks
    
    def get_document_summary(self, document_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a summary of the document
        
        Args:
            document_data (Dict[str, Any]): Document data from process_document
            
        Returns:
            Dict[str, Any]: Document summary
        """
        content = document_data['content']
        
        # Basic statistics
        words = content.split()
        sentences = content.split('.')
        paragraphs = content.split('\n\n')
        
        # Get first few sentences as preview
        preview = '. '.join(sentences[:3]) + '.' if len(sentences) > 3 else content[:200] + "..."
        
        return {
            'file_name': document_data['file_name'],
            'file_type': document_data['file_type'],
            'file_size_mb': round(document_data['file_size'] / (1024 * 1024), 2),
            'word_count': document_data['word_count'],
            'char_count': document_data['char_count'],
            'sentence_count': len([s for s in sentences if s.strip()]),
            'paragraph_count': len([p for p in paragraphs if p.strip()]),
            'preview': preview,
            'processed_at': document_data.get('processed_at', 'Unknown')
        }

