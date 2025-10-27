"""
User router for user-specific functions
"""
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from sqlalchemy.orm import Session
from typing import List, Optional
from database import get_db
from models import User, Test, Document, Result, Question
from schemas import (
    UserResponse, TestResponse, DocumentResponse, QuizSubmission, QuizResult
)
from utils.auth import get_current_student, get_current_user

router = APIRouter()

@router.get("/stats")
async def get_user_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get user dashboard statistics (works for both admin and student)"""
    # Get available tests count
    total_tests = db.query(Test).filter(Test.is_active == True).count()
    
    # Get completed tests count
    completed_tests = db.query(Result).filter(Result.user_id == current_user.id).count()
    
    # Get average score
    results = db.query(Result).filter(Result.user_id == current_user.id).all()
    average_score = sum([r.score for r in results]) / len(results) if results else 0
    
    # Get documents count
    total_documents = db.query(Document).filter(Document.user_id == current_user.id).count()
    
    return {
        "total_tests": total_tests,
        "completed_tests": completed_tests,
        "average_score": round(average_score, 1),
        "total_documents": total_documents
    }

@router.get("/profile", response_model=UserResponse)
async def get_user_profile(
    current_user: User = Depends(get_current_student)
):
    """Get current user profile"""
    return UserResponse(
        id=current_user.id,
        username=current_user.username,
        email=current_user.email,
        role=current_user.role.value,
        test_history=current_user.test_history or [],
        created_at=current_user.created_at
    )

@router.get("/tests", response_model=List[TestResponse])
async def get_available_tests(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_student)
):
    """Get all available tests for the user"""
    tests = db.query(Test).filter(Test.is_active == True).all()
    return tests

@router.get("/tests/{test_id}", response_model=TestResponse)
async def get_test_for_user(
    test_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_student)
):
    """Get test details for user (without correct answers)"""
    test = db.query(Test).filter(
        Test.id == test_id,
        Test.is_active == True
    ).first()
    
    if not test:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Test not found"
        )
    
    return test

@router.post("/tests/{test_id}/submit", response_model=QuizResult)
async def submit_test(
    test_id: int,
    submission: QuizSubmission,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_student)
):
    """Submit test answers and get results"""
    # Verify test exists and is active
    test = db.query(Test).filter(
        Test.id == test_id,
        Test.is_active == True
    ).first()
    
    if not test:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Test not found"
        )
    
    # Get test questions with correct answers
    questions = db.query(Question).filter(Question.test_id == test_id).all()
    
    # Calculate score
    correct_answers = 0
    total_questions = len(questions)
    user_answers = {}
    
    for answer in submission.answers:
        question = next((q for q in questions if q.id == answer.question_id), None)
        if question:
            user_answers[str(answer.question_id)] = answer.answer
            if answer.answer == question.correct_answer:
                correct_answers += 1
    
    score = (correct_answers / total_questions) * 100 if total_questions > 0 else 0
    
    # Create result record
    result = Result(
        user_id=current_user.id,
        test_id=test_id,
        score=score,
        total_questions=total_questions,
        correct_answers=correct_answers,
        time_taken_minutes=submission.time_taken_minutes,
        answers=user_answers
    )
    
    db.add(result)
    db.commit()
    db.refresh(result)
    
    return result

@router.get("/results", response_model=List[QuizResult])
async def get_user_results(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_student)
):
    """Get user's test results"""
    results = db.query(Result).filter(Result.user_id == current_user.id).all()
    return results

@router.get("/results/{result_id}", response_model=QuizResult)
async def get_result_details(
    result_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_student)
):
    """Get specific result details"""
    result = db.query(Result).filter(
        Result.id == result_id,
        Result.user_id == current_user.id
    ).first()
    
    if not result:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Result not found"
        )
    
    return result

@router.get("/documents", response_model=List[DocumentResponse])
async def get_user_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get user's uploaded documents"""
    documents = db.query(Document).filter(Document.user_id == current_user.id).all()
    return documents

@router.post("/documents")
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Upload a document with proper metadata extraction"""
    import os
    from datetime import datetime
    import shutil
    
    try:
        print(f"\n📤 ===== UPLOAD REQUEST =====")
        print(f"User ID: {current_user.id}")
        print(f"Filename: {file.filename}")
        
        # 🔒 PREVENT DUPLICATE UPLOADS
        existing_doc = db.query(Document).filter(
            Document.user_id == current_user.id,
            Document.doc_name == file.filename
        ).first()
        
        if existing_doc:
            print(f"⚠️  Document already exists: {file.filename}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document with this name already uploaded. Please rename or delete the existing one."
            )
        
        # Create uploads directory
        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{timestamp}_{file.filename}"
        file_path = os.path.join(upload_dir, filename)
        
        # Save file
        print(f"💾 Saving to: {file_path}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Get file size in bytes
        file_size = os.path.getsize(file_path)
        print(f"✅ File saved: {file_size} bytes")
        
        # Determine file type
        file_extension = os.path.splitext(file.filename)[1].lower()
        file_type = file_extension[1:] if file_extension else 'unknown'
        
        # Extract metadata (pages & words)
        total_pages = 0
        total_words = 0
        
        try:
            if file_type == 'pdf':
                import fitz  # PyMuPDF
                pdf = fitz.open(file_path)
                total_pages = len(pdf)
                total_words = sum(len(page.get_text().split()) for page in pdf)
                pdf.close()
                print(f"📄 PDF: {total_pages} pages, {total_words} words")
                
            elif file_type in ['doc', 'docx']:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                total_words = sum(len(p.text.split()) for p in doc.paragraphs)
                total_pages = max(1, len(doc.paragraphs) // 30)
                print(f"📝 DOCX: ~{total_pages} pages, {total_words} words")
                
            elif file_type == 'txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                total_words = len(text.split())
                total_pages = max(1, len(text) // 2000)
                print(f"📃 TXT: ~{total_pages} pages, {total_words} words")
                
        except Exception as e:
            print(f"⚠️  Metadata extraction failed: {e}")
            # Continue anyway with defaults
        
        # Create document record with all metadata
        document = Document(
            user_id=current_user.id,
            doc_name=file.filename,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            total_words=total_words,
            total_pages=total_pages,
            is_processed=True
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        print(f"✅ Upload complete: ID={document.id}")
        print(f"===== END UPLOAD =====\n")
        
        # Return response
        return {
            "id": document.id,
            "doc_name": document.doc_name,
            "file_type": document.file_type,
            "file_size": document.file_size,
            "total_words": document.total_words,
            "total_pages": document.total_pages,
            "is_processed": document.is_processed,
            "created_at": document.created_at.isoformat() if document.created_at else None,
            "user_id": document.user_id,
            "file_path": document.file_path
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ ===== UPLOAD FAILED =====")
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        print(f"===== END ERROR =====\n")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Upload failed: {str(e)}"
        )

@router.post("/documents/{document_id}/process")
async def process_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_student)
):
    """Manually trigger document processing"""
    from utils.document_processor import DocumentProcessor
    
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        doc_processor = DocumentProcessor()
        result = doc_processor.process_document(document.file_path, str(document.id))
        
        if result.get("success"):
            text_content = result.get("text_content", {})
            document.total_words = text_content.get("total_words", 0)
            document.total_pages = text_content.get("total_pages", 0)
            document.is_processed = True
            
            db.commit()
            db.refresh(document)
            
            return {
                "message": "Document processed successfully",
                "total_words": document.total_words,
                "total_pages": document.total_pages
            }
        else:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Processing failed: {result.get('error', 'Unknown error')}"
            )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing document: {str(e)}"
        )

@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_student)
):
    """Delete user's document"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Delete file from filesystem
    import os
    if os.path.exists(document.file_path):
        os.remove(document.file_path)
    
    db.delete(document)
    db.commit()
    
    return {"message": "Document deleted successfully"}
