"""
Utility functions for extracting text from various document formats
"""
import PyPDF2
import docx
import os

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        
        return "\n".join(text_content)
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        doc = docx.Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return "\n".join(text_content)
    except Exception as e:
        raise Exception(f"Error extracting DOCX text: {str(e)}")

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        raise Exception(f"Error extracting TXT text: {str(e)}")

def get_pdf_page_count(file_path: str) -> int:
    """Get the number of pages in a PDF"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            return len(pdf_reader.pages)
    except:
        return 0

def get_docx_page_count(file_path: str) -> int:
    """Estimate the number of pages in a DOCX (rough estimate)"""
    try:
        doc = docx.Document(file_path)
        # Rough estimate: count paragraphs and divide by ~25 paragraphs per page
        return max(1, len(doc.paragraphs) // 25)
    except:
        return 1

